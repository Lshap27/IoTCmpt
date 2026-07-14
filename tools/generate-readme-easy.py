from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "README_easy_source.md"
OUTPUT = ROOT / "README_easy.docx"

# compact_reference_guide preset, with a named CJK fallback override.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "667085"
PALE_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
GOLD = "B7791F"
INK = "182230"


def set_run_font(
    run, size: float | None = None, bold: bool | None = None, color: str | None = None
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, 9, color=MUTED)


def configure_header_footer(section) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("IoTCmpt  |  队友答辩与演示手册")
    set_run_font(run, 8.5, bold=True, color=MUTED)
    p_pr = header._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D6DEE8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    label = footer.add_run("软件验收版  ·  第 ")
    set_run_font(label, 9, color=MUTED)
    add_page_field(footer)
    suffix = footer.add_run(" 页")
    set_run_font(suffix, 9, color=MUTED)


def add_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    next_id = max(existing, default=-1) + 1

    def create(fmt: str, text: str, font: str | None = None) -> int:
        nonlocal next_id
        abstract_id = next_id
        next_id += 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.extend([tabs, ind])
        level.extend([start, num_fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            level.append(r_pr)
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num_id = (
            max(
                [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))],
                default=0,
            )
            + 1
        )
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)
        return num_id

    return create("bullet", "•"), create("decimal", "%1.")


def apply_list(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def clone_numbering_instance(doc: Document, source_num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    source = next(
        item
        for item in numbering.findall(qn("w:num"))
        if int(item.get(qn("w:numId"))) == source_num_id
    )
    abstract_id = source.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_id = (
        max(int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))) + 1
    )
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), abstract_id)
    num.append(ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return new_id


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, 9.5, color=DARK_BLUE)
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True, color=NAVY)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_callout(doc: Document, text: str, kind: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.06)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(9)
    p_pr = paragraph._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "FFF7E6" if kind == "IMPORTANT" else CALLOUT)
    p_pr.append(shade)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), GOLD if kind == "IMPORTANT" else BLUE)
    borders.append(left)
    p_pr.append(borders)
    label = paragraph.add_run("重点  " if kind == "IMPORTANT" else "说明  ")
    set_run_font(label, bold=True, color=GOLD if kind == "IMPORTANT" else BLUE)
    add_inline(paragraph, text)


def add_image(doc: Document, alt: str, relative: str) -> None:
    image_path = ROOT / relative.replace("/", "\\")
    if not image_path.exists():
        add_callout(
            doc,
            f"截图未找到：{relative}。重新生成手册前请启动演示栈并采集界面。",
            "NOTE",
        )
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.25))
    if image_path.name == "startup-panel.png":
        shape.height = Inches(2.75)
        blip_fill = shape._inline.graphic.graphicData.pic.blipFill
        src_rect = OxmlElement("a:srcRect")
        src_rect.set("t", "25000")
        blip_fill.insert(1, src_rect)
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(f"图：{alt}（本机地址与秘密不进入文档）")
    set_run_font(run, 9, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("AIoT 架构 · 可靠性 · 演示 · 上板")
    set_run_font(run, 11, bold=True, color=GOLD)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("IoTCmpt 队友答辩与演示手册")
    set_run_font(run, 29, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("不要求编程基础，也能讲清系统怎样工作、为什么可靠、怎样扩展")
    set_run_font(run, 14, color=DARK_BLUE)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.6)
    lead.paragraph_format.right_indent = Inches(0.6)
    lead.paragraph_format.space_after = Pt(54)
    run = lead.add_run(
        "现场安全由固件保证；可靠协作由 Gateway 与数据库保证；智能分析由独立 Worker + MCP 保证；用户操作与排障由双面板保证。"
    )
    set_run_font(run, 12, bold=True, color=BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run("架构 v2 · 软件验收版")
    set_run_font(run, 11, bold=True, color=NAVY)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("2026 年 7 月 15 日")
    set_run_font(run, 10.5, color=MUTED)


def render_segment(
    doc: Document, segment: str, bullet_id: int, decimal_id: int
) -> None:
    for raw in segment.strip().splitlines():
        line = raw.strip()
        if not line:
            continue
        image_match = re.fullmatch(r"!\[([^]]+)]\(([^)]+)\)", line)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            continue
        if line.startswith("> [!IMPORTANT]"):
            add_callout(doc, line.removeprefix("> [!IMPORTANT]").strip(), "IMPORTANT")
            continue
        if line.startswith("> [!NOTE]"):
            add_callout(doc, line.removeprefix("> [!NOTE]").strip(), "NOTE")
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline(paragraph, line[4:])
            continue
        if line.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            heading_text = line[3:]
            numbered = re.match(r"^(\d+\.)\s+(.*)$", heading_text)
            if numbered:
                number_run = paragraph.add_run(numbered.group(1) + "\u00a0")
                set_run_font(number_run, bold=True, color=BLUE)
                add_inline(paragraph, numbered.group(2))
            else:
                add_inline(paragraph, heading_text)
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph()
            apply_list(paragraph, bullet_id)
            add_inline(paragraph, line[2:])
            continue
        if re.match(r"^\d+\. ", line):
            paragraph = doc.add_paragraph()
            apply_list(paragraph, decimal_id)
            add_inline(paragraph, re.sub(r"^\d+\. ", "", line))
            continue
        paragraph = doc.add_paragraph()
        add_inline(paragraph, line)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    configure_page(section)
    configure_header_footer(section)
    configure_styles(doc)
    bullet_id, decimal_id = add_numbering(doc)
    doc.core_properties.title = "IoTCmpt 队友答辩与演示手册"
    doc.core_properties.subject = "IoTCmpt 架构、运行逻辑、演示、上板与排障"
    doc.core_properties.author = "IoTCmpt 项目组"
    doc.core_properties.keywords = "IoTCmpt, AIoT, ESP32-S3, MCP, MQTT, 答辩"

    add_cover(doc)
    segments = re.split(
        r"(?m)^\[\[PAGEBREAK\]\]\s*$", SOURCE.read_text(encoding="utf-8")
    )[1:]
    for segment in segments:
        doc.add_page_break()
        render_segment(
            doc,
            segment,
            clone_numbering_instance(doc, bullet_id),
            clone_numbering_instance(doc, decimal_id),
        )

    # Keep one portrait section and explicit page geometry throughout.
    if len(doc.sections) != 1:
        for extra in doc.sections[1:]:
            extra.start_type = WD_SECTION.NEW_PAGE
            configure_page(extra)
            configure_header_footer(extra)
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT} with {1 + len(segments)} planned pages")


if __name__ == "__main__":
    build()
